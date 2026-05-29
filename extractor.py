"""
extractor.py
Tarefa 1 – Extração de Texto Multi-formato
Suporta PDF, DOCX e TXT.
"""

import io
import os


def extrair_txt(conteudo: bytes) -> str:
    """
    Lê ficheiros TXT detetando o encoding automaticamente com chardet.
    Se chardet não estiver instalado, usa UTF-8 por defeito.
    """
    try:
        import chardet
        enc = chardet.detect(conteudo).get("encoding") or "utf-8"
    except ImportError:
        enc = "utf-8"
    try:
        return conteudo.decode(enc)
    except Exception:
        # se ainda assim falhar, substitui os bytes inválidos por ?
        return conteudo.decode("utf-8", errors="replace")


def extrair_pdf(conteudo: bytes) -> str:
    """
    Extrai texto de PDF usando pdfplumber.
    O \f (form feed) é inserido entre páginas como separador
    para ajudar a detetar cabeçalhos/rodapés repetidos.
    """
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("Instala pdfplumber:  pip install pdfplumber")

    paginas = []
    with pdfplumber.open(io.BytesIO(conteudo)) as pdf:
        for pagina in pdf.pages:
            texto = pagina.extract_text(x_tolerance=2, y_tolerance=2)
            if texto:
                paginas.append(texto)
            paginas.append("\f")  # separador de página
    return "\n".join(paginas)


def extrair_docx(conteudo: bytes) -> str:
    """
    Extrai texto de ficheiros Word (.docx) usando python-docx.
    Extrai tanto os parágrafos normais como o texto dentro de tabelas.
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Instala python-docx:  pip install python-docx")

    doc = Document(io.BytesIO(conteudo))
    linhas = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            linhas.append("\t".join(c.text for c in linha.cells))
    return "\n".join(linhas)


def extrair_texto(caminho: str) -> str:
    """
    Ponto de entrada da extração.
    Lê o ficheiro em modo binário e chama a função correta
    com base na extensão do ficheiro.
    """
    if not os.path.isfile(caminho):
        raise FileNotFoundError(f"Ficheiro não encontrado: {caminho}")

    with open(caminho, "rb") as f:
        conteudo = f.read()

    nome = caminho.lower()
    if nome.endswith(".pdf"):
        return extrair_pdf(conteudo)
    elif nome.endswith(".docx"):
        return extrair_docx(conteudo)
    elif nome.endswith(".txt"):
        return extrair_txt(conteudo)
    else:
        raise ValueError("Formato não suportado. Use PDF, DOCX ou TXT.")