# ── Bibliotecas da biblioteca padrão do Python (não precisam de ser instaladas) ──
import io           # permite trabalhar com ficheiros em memória (sem guardar no disco)
import os           # acesso ao sistema de ficheiros (verificar se ficheiro existe, etc.)
import re           # expressões regulares — para encontrar e substituir padrões no texto
import json         # converter entre texto JSON e dicionários Python
import time         # medir o tempo de execução (ex: quanto demora a API a responder)
import html as html_module   # escapar caracteres especiais em HTML (ex: < vira &lt;)
import datetime     # obter a data/hora atual para o relatório
import unicodedata  # normalização de caracteres Unicode
import urllib.request   # fazer pedidos HTTP sem bibliotecas externas
import urllib.error     # tratar erros de rede (timeout, 404, etc.)
import threading    # executar código em paralelo — usado para não bloquear a interface durante a API
from collections import Counter  # contar frequências (ex: quantas vezes uma linha aparece)

# ── Biblioteca de interface gráfica (já incluída no Python) ──
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
# ttk          → widgets com melhor aspeto visual (tabs, progress bar, combobox)
# filedialog   → janela de "Abrir ficheiro" / "Guardar como"
# messagebox   → janelas de aviso/erro/sucesso
# scrolledtext → caixa de texto com barra de scroll


# =============================================================================
# TAREFA 1 – EXTRAÇÃO DE TEXTO MULTI-FORMATO
# =============================================================================

def extrair_txt(conteudo: bytes) -> str:
    """
    Lê um ficheiro TXT em bruto.
    O 'conteudo' é o ficheiro em bytes — temos de o converter para texto.
    Usamos chardet para detetar automaticamente o encoding (UTF-8, Latin-1, etc.)
    porque nem todos os ficheiros TXT usam o mesmo sistema de codificação.
    """
    try:
        import chardet
        # chardet analisa os bytes e tenta adivinhar o encoding
        enc = chardet.detect(conteudo).get("encoding") or "utf-8"
    except ImportError:
        # se chardet não estiver instalado, assume UTF-8 por defeito
        enc = "utf-8"
    try:
        return conteudo.decode(enc)
    except Exception:
        # se ainda assim falhar, usa UTF-8 e substitui os caracteres inválidos por ?
        return conteudo.decode("utf-8", errors="replace")


def extrair_pdf(conteudo: bytes) -> str:
    """
    Extrai o texto de um ficheiro PDF usando a biblioteca pdfplumber.
    Percorre todas as páginas do PDF e junta o texto de cada uma.
    O \f (form feed) é inserido entre páginas como separador — é útil
    para mais tarde detetar cabeçalhos/rodapés repetidos entre páginas.
    x_tolerance e y_tolerance controlam a precisão da extração do texto.
    """
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("Instala pdfplumber:  pip install pdfplumber")

    paginas = []
    # io.BytesIO converte os bytes do ficheiro num "ficheiro em memória"
    # para que o pdfplumber consiga abrir sem precisar de um ficheiro real
    with pdfplumber.open(io.BytesIO(conteudo)) as pdf:
        for pagina in pdf.pages:
            texto = pagina.extract_text(x_tolerance=2, y_tolerance=2)
            if texto:
                paginas.append(texto)
            paginas.append("\f")   # marcador de fim de página
    return "\n".join(paginas)


def extrair_docx(conteudo: bytes) -> str:
    """
    Extrai o texto de um ficheiro Word (.docx) usando python-docx.
    Um DOCX é um ficheiro ZIP com XML dentro — o python-docx trata disso por nós.
    Extrai tanto os parágrafos normais como o texto dentro de tabelas.
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Instala python-docx:  pip install python-docx")

    doc = Document(io.BytesIO(conteudo))
    # extrai o texto de cada parágrafo
    linhas = [p.text for p in doc.paragraphs]
    # extrai também o texto das tabelas (linha a linha, célula a célula)
    for tabela in doc.tables:
        for linha in tabela.rows:
            linhas.append("\t".join(c.text for c in linha.cells))
    return "\n".join(linhas)


def extrair_texto(caminho: str) -> str:
    """
    Ponto de entrada da extração — a função que chamamos na interface.
    Recebe o caminho do ficheiro, lê-o em modo binário (rb = read binary)
    e decide qual das três funções acima usar com base na extensão.
    Lemos sempre em binário primeiro para preservar tudo — a conversão
    para texto é feita dentro de cada função específica.
    """
    if not os.path.isfile(caminho):
        raise FileNotFoundError(f"Ficheiro não encontrado: {caminho}")

    # abre o ficheiro em modo binário para preservar todos os bytes originais
    with open(caminho, "rb") as f:
        conteudo = f.read()

    nome = caminho.lower()
    if nome.endswith(".pdf"):
        return extrair_pdf(conteudo)
    elif nome.endswith(".docx"):
        return extrair_docx(conteudo)
    elif nome.endswith(".txt"):
        return extrair_txt(conteudo)
    else:
        raise ValueError(f"Formato não suportado. Use PDF, DOCX ou TXT.")


# =============================================================================
# TAREFA 2 – PIPELINE DE LIMPEZA E PRÉ-PROCESSAMENTO
# =============================================================================

def remover_artefactos(texto: str) -> str:
    """
    Remove "lixo" do texto — caracteres que não são texto real.
    [\x00-\x08\x0b\x0c\x0e-\x1f\x7f] são caracteres de controlo ASCII
    (codigos abaixo de 32) que não têm representação visual — surgem
    frequentemente em PDFs e documentos digitalizados.
    \ufffd é o caractere de substituição que aparece quando há erros de encoding.
    Os seguintes são caracteres invisíveis de largura zero comuns em copy-paste
    de páginas web: zero-width space, zero-width joiner, soft hyphen, etc.
    """
    # remove caracteres de controlo (exceto \n newline e \t tab que são úteis)
    texto = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", texto)
    # remove o caractere de substituição (sinal de encoding corrompido)
    texto = texto.replace("\ufffd", "")
    # remove caracteres invisíveis de largura zero
    texto = re.sub(r"[\u200b\u200c\u200d\ufeff\u00ad]", "", texto)
    return texto


def normalizar_unicode(texto: str) -> str:
    """
    Normaliza o texto para a forma NFC (Canonical Decomposition followed by
    Canonical Composition). O mesmo caractere pode ser representado de formas
    diferentes em Unicode — por exemplo, 'é' pode ser um único caractere (NFC)
    ou dois caracteres separados: 'e' + acento (NFD). Esta função garante
    consistência, especialmente entre ficheiros criados em Mac vs Windows.
    """
    return unicodedata.normalize("NFC", texto)


def corrigir_quebras_de_linha(texto: str) -> str:
    """
    Normaliza as quebras de linha para sempre usar \n (estilo Unix/Linux).
    Windows usa \r\n (CRLF), Mac antigo usava \r (CR) — esta função unifica tudo.
    O \f (form feed / nova página) é convertido em duas linhas vazias,
    que é uma separação de parágrafo normal.
    """
    texto = texto.replace("\r\n", "\n")   # Windows → Unix
    texto = texto.replace("\r", "\n")      # Mac antigo → Unix
    texto = texto.replace("\f", "\n\n")    # nova página → linha em branco
    return texto


def remover_numeros_pagina(texto: str) -> str:
    """
    Remove linhas que sejam apenas números de página.
    Em PDFs, os números de página ficam frequentemente como linhas sozinhas.
    Usa expressões regulares para detetar padrões como:
    - "3", "- 3 -", "– 3 –" (número sozinho numa linha)
    - "Page 3", "Página 3", "Pág. 3" (com prefixo de página)
    re.fullmatch exige que o padrão ocupe a linha TODA (não apenas parte dela).
    """
    linhas = texto.split("\n")
    resultado = []
    for linha in linhas:
        s = linha.strip()   # remove espaços no início e fim para comparar
        # padrão: apenas dígitos, opcionalmente rodeados de traços/espaços
        if re.fullmatch(r"[-–\s]*\d+[-–\s]*", s):
            continue   # salta esta linha (é número de página)
        # padrão: "Page X" ou "Página X" no início da linha
        if re.match(r"^(page|página|pág\.?)\s*\d+", s, re.IGNORECASE):
            continue
        resultado.append(linha)
    return "\n".join(resultado)


def remover_cabecalhos_rodapes(texto: str) -> str:
    """
    Deteta e remove cabeçalhos e rodapés repetidos ao longo do documento.
    A lógica é: se uma linha curta (menos de 120 chars) aparece em pelo menos
    40% das páginas, provavelmente é um cabeçalho ou rodapé.
    Usamos Counter para contar quantas vezes cada linha aparece.
    O threshold (mínimo de repetições) depende do número de páginas.
    """
    linhas = texto.split("\n")
    # conta quantas páginas tem o documento (pelo número de \f que colocámos na extração)
    n_paginas = texto.count("\f") + 1
    # uma linha precisa de aparecer em pelo menos 40% das páginas para ser considerada cabeçalho
    threshold = max(2, n_paginas * 0.4)
    # conta a frequência de cada linha não vazia
    freq = Counter(l.strip() for l in linhas if l.strip())
    # mantém apenas as linhas que NÃO são cabeçalhos/rodapés
    return "\n".join(
        l for l in linhas
        if not (freq.get(l.strip(), 0) >= threshold and len(l.strip()) < 120)
    )


def reconstruir_paragrafos(texto: str) -> str:
    """
    Re-une linhas que foram partidas a meio de uma frase — problema muito
    comum em PDFs, onde o texto é extraído linha a linha da página física.
    Por exemplo, um parágrafo que na página ocupa 3 linhas físicas fica
    extraído como 3 linhas separadas, quando devia ser uma frase contínua.
    
    A lógica: se a linha atual tem mais de 40 chars (não é um título curto)
    e NÃO termina em pontuação final (. ! ? : ; - —), então é provável que
    a frase continue na próxima linha — juntamo-las com um espaço.
    Se a próxima linha tem indentação (2+ espaços), é um novo parágrafo.
    """
    linhas = texto.split("\n")
    resultado = []
    i = 0
    while i < len(linhas):
        atual = linhas[i].rstrip()
        proxima = linhas[i + 1].strip() if i + 1 < len(linhas) else ""
        # condições para juntar com a linha seguinte:
        if (atual                           # linha atual não está vazia
                and proxima                 # linha seguinte não está vazia
                and len(atual) > 40         # linha longa o suficiente (não é título)
                and not atual.endswith((".", "!", "?", ":", ";", "-", "—"))  # não termina frase
                and not re.match(r"^\s{2,}", linhas[i + 1])):  # próxima não tem indentação
            resultado.append(atual + " " + proxima)
            i += 2   # salta a próxima linha (já foi incorporada)
        else:
            resultado.append(atual)
            i += 1
    return "\n".join(resultado)


def normalizar_pontuacao(texto: str) -> str:
    """
    Converte caracteres tipográficos "bonitos" para os equivalentes ASCII simples.
    Processadores de texto (Word, Pages) e websites usam versões especiais:
    - aspas "curvas" (" " ' ') → aspas retas (" ')
    - travessão longo (— –) → hífen simples (-)
    - reticências (…) → três pontos (...)
    - bullet (•) → hífen (-)
    Fazemos isto porque modelos de linguagem e ferramentas NLP funcionam melhor
    com texto ASCII limpo.
    """
    substituicoes = {
        "\u2018": "'",   # aspa simples esquerda
        "\u2019": "'",   # aspa simples direita
        "\u201c": '"',   # aspa dupla esquerda
        "\u201d": '"',   # aspa dupla direita
        "\u2013": "-",   # en-dash
        "\u2014": "-",   # em-dash (travessão longo)
        "\u2026": "...", # reticências
        "\u2022": "-",   # bullet point
    }
    for orig, repl in substituicoes.items():
        texto = texto.replace(orig, repl)
    return texto


def normalizar_espacos(texto: str) -> str:
    """
    Limpa espaços em excesso no texto.
    [ \t]+ → colapsa múltiplos espaços/tabs numa linha para um só espaço
    \n{3,} → substitui 3 ou mais linhas em branco seguidas por só 2
    rstrip() em cada linha → remove espaços desnecessários no final de cada linha
    No final, strip() remove linhas em branco no início e fim do texto todo.
    """
    texto = re.sub(r"[ \t]+", " ", texto)             # múltiplos espaços → um
    texto = re.sub(r"\n{3,}", "\n\n", texto)           # muitas linhas vazias → duas
    texto = "\n".join(l.rstrip() for l in texto.split("\n"))  # limpa fim de cada linha
    return texto.strip()


# Lista ordenada de todas as etapas da pipeline — (nome para mostrar, função a chamar)
# A ORDEM IMPORTA: por exemplo, corrigir quebras de linha antes de reconstruir parágrafos
ETAPAS_PIPELINE = [
    ("Remoção de artefactos",        remover_artefactos),
    ("Normalização Unicode",          normalizar_unicode),
    ("Correção de quebras de linha",  corrigir_quebras_de_linha),
    ("Remoção de números de página",  remover_numeros_pagina),
    ("Deteção cabeçalhos/rodapés",    remover_cabecalhos_rodapes),
    ("Reconstrução de parágrafos",    reconstruir_paragrafos),
    ("Normalização de pontuação",     normalizar_pontuacao),
    ("Normalização de espaços",       normalizar_espacos),
]


def correr_pipeline(texto: str, etapas_ativas: list) -> tuple:
    """
    Executa a pipeline de limpeza — aplica as etapas uma a uma, pela ordem da lista.
    'etapas_ativas' é uma lista de True/False com a mesma dimensão que ETAPAS_PIPELINE,
    que diz quais as etapas que o utilizador ativou nos checkboxes da interface.
    Para cada etapa executada, guarda estatísticas (chars antes, depois, diferença).
    Devolve o texto final limpo e a lista de estatísticas para mostrar na interface.
    """
    estatisticas = []
    texto_atual = texto

    for (nome, funcao), ativo in zip(ETAPAS_PIPELINE, etapas_ativas):
        if not ativo:
            continue   # utilizador desativou esta etapa — salta

        antes = texto_atual        # guarda o texto antes desta etapa
        texto_atual = funcao(texto_atual)   # aplica a etapa

        # guarda as estatísticas desta etapa para mostrar na interface
        estatisticas.append({
            "etapa":           nome,
            "chars_antes":     len(antes),
            "chars_depois":    len(texto_atual),
            "chars_removidos": len(antes) - len(texto_atual),   # pode ser negativo
            "linhas_antes":    antes.count("\n") + 1,
            "linhas_depois":   texto_atual.count("\n") + 1,
        })

    return texto_atual, estatisticas


# =============================================================================
# TAREFA 3 – PREPARAÇÃO DO INPUT PARA O SLM
# =============================================================================

# ── Deteção de idioma ─────────────────────────────────────────────────────────

# Perfis de trigramas por idioma — sequências de 3 caracteres mais comuns
# em cada língua. Por exemplo, "de " e " o " são muito mais comuns em português
# do que em inglês. "the" e "ing" são típicos do inglês.
PERFIS_IDIOMA = {
    "pt": ["de ", " de", " a ", "que", " o ", " e ", "ão", "os ", "por", "não", "para", "uma", "com"],
    "en": ["the", " th", "he ", "in ", " of", "and", " an", "to ", "is ", "it ", "for"],
    "es": ["de ", " de", " la", "la ", "que", " en", " y ", " el", "con", "una", "ión"],
    "fr": ["de ", " de", " la", "le ", "les", " et", "des", "une", " à ", "nt "],
}

NOMES_IDIOMA = {
    "pt": "Português", "en": "English",
    "es": "Español",   "fr": "Français",
    "desconhecido": "Desconhecido"
}


def detetar_idioma(texto: str) -> tuple:
    """
    Deteta o idioma do texto por análise de trigramas, sem bibliotecas externas.
    Um trigrama é uma sequência de 3 caracteres consecutivos.
    O algoritmo:
    1. Pega nos primeiros 3000 caracteres do texto (amostra suficiente)
    2. Conta todos os trigramas presentes
    3. Para cada idioma, soma as frequências dos seus trigramas característicos
    4. O idioma com maior pontuação é o vencedor
    A confiança é calculada pela margem entre o 1º e o 2º classificado —
    se for muito próximo, significa que não temos a certeza.
    Devolve (código_idioma, confiança entre 0 e 1).
    """
    if not texto or len(texto) < 50:
        return "desconhecido", 0.0

    # normaliza e limita o texto para análise rápida
    amostra = re.sub(r"\s+", " ", texto.lower()[:3000])

    # conta todos os trigramas (sequências de 3 chars) na amostra
    trigramas = Counter(amostra[i:i+3] for i in range(len(amostra) - 2))
    total = sum(trigramas.values()) or 1   # evitar divisão por zero

    # calcula pontuação para cada idioma
    pontuacoes = {
        idioma: sum(trigramas.get(ng, 0) for ng in perfil) / total
        for idioma, perfil in PERFIS_IDIOMA.items()
    }

    melhor = max(pontuacoes, key=pontuacoes.get)
    valores = sorted(pontuacoes.values(), reverse=True)

    # margem entre 1º e 2º — quanto maior, mais certos estamos
    margem = valores[0] - valores[1] if len(valores) > 1 else valores[0]
    confianca = min(1.0, margem * 20)   # escala a margem para [0, 1]

    if confianca < 0.1:
        return "desconhecido", round(confianca, 3)
    return melhor, round(confianca, 3)


# ── Chunking (segmentação em blocos) ─────────────────────────────────────────

def chunk_tamanho_fixo(texto: str, tamanho: int = 512, overlap: int = 50) -> list:
    """
    Divide o texto em blocos de tamanho fixo (em caracteres).
    O 'overlap' é a sobreposição entre chunks consecutivos — os últimos
    50 chars de um chunk são os primeiros do próximo. Isto é importante
    para que o modelo não perca contexto nas fronteiras entre chunks.
    Tenta cortar sempre num espaço (rfind(" ")) para não partir palavras a meio.
    """
    chunks = []
    inicio = 0
    while inicio < len(texto):
        fim = min(inicio + tamanho, len(texto))
        # tenta encontrar o último espaço antes do limite para não partir palavras
        if fim < len(texto):
            corte = texto.rfind(" ", inicio, fim)
            if corte > inicio:
                fim = corte
        bloco = texto[inicio:fim].strip()
        if bloco:
            chunks.append(bloco)
        # o próximo chunk começa 'overlap' chars antes do fim deste
        inicio = fim - overlap if fim < len(texto) else fim
    return chunks


def chunk_por_paragrafo(texto: str, max_chars: int = 1024) -> list:
    """
    Divide o texto em blocos por parágrafos (separados por linha em branco).
    Agrupa parágrafos consecutivos enquanto o total não ultrapassar max_chars.
    Quando um parágrafo novo faria o bloco ultrapassar o limite, começa
    um novo chunk. Garante que os chunks têm contexto semântico coerente.
    """
    chunks, buffer = [], ""
    # re.split com \n{2,} separa nos parágrafos (2+ linhas em branco)
    for para in re.split(r"\n{2,}", texto):
        para = para.strip()
        if not para:
            continue
        # se o parágrafo cabe no buffer atual, adiciona
        if len(buffer) + len(para) + 2 <= max_chars:
            buffer = (buffer + "\n\n" + para).strip() if buffer else para
        else:
            # o buffer está cheio — guarda-o e começa novo com este parágrafo
            if buffer:
                chunks.append(buffer)
            buffer = para
    if buffer:   # não esquecer o último buffer
        chunks.append(buffer)
    return chunks


def chunk_por_frase(texto: str, max_chars: int = 512) -> list:
    """
    Divide o texto em blocos por frases (terminadas em . ! ?).
    O padrão regex (?<=[.!?]) segue-se de espaco e (?=[A-Z...]) usa lookbehind e lookahead:
    - (?<=[.!?]) → antes tem de haver pontuação final
    - (?=[A-Z]) → depois tem de começar com maiúscula (nova frase)
    Agrupa frases até atingir max_chars, mantendo coerência semântica.
    """
    chunks, buffer = [], ""
    frases = re.split(r"(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚÀÂÊÔÃÕ\"])", texto)
    for frase in frases:
        frase = frase.strip()
        if not frase:
            continue
        if len(buffer) + len(frase) + 1 <= max_chars:
            buffer = (buffer + " " + frase).strip() if buffer else frase
        else:
            if buffer:
                chunks.append(buffer)
            buffer = frase
    if buffer:
        chunks.append(buffer)
    return chunks


# Dicionário com todas as estratégias disponíveis — nome → função
# Usado para preencher o dropdown da interface e chamar a função certa
ESTRATEGIAS_CHUNKING = {
    "Tamanho fixo (512 chars)":  lambda t: chunk_tamanho_fixo(t, 512, 50),
    "Tamanho fixo (1024 chars)": lambda t: chunk_tamanho_fixo(t, 1024, 100),
    "Por parágrafo":             lambda t: chunk_por_paragrafo(t, 1024),
    "Por frase":                 lambda t: chunk_por_frase(t, 512),
}


# ── Geração automática de prompts ─────────────────────────────────────────────

# Templates de prompts por idioma e tipo de normalização.
# Cada template tem {texto} como placeholder que é substituído pelo chunk real.
# Ter prompts no idioma do texto melhora significativamente a qualidade
# da resposta do modelo — um modelo responde melhor se for "perguntado" no mesmo idioma.
TEMPLATES_PROMPT = {
    "pt": {
        "Normalização geral":  "Normaliza o seguinte texto em português, corrigindo erros ortográficos, melhorando a pontuação e a coesão textual. Mantém o significado original e devolve apenas o texto normalizado:\n\n{texto}",
        "Registo formal":      "Reformula o seguinte texto em português num registo formal e profissional, corrigindo erros. Devolve apenas o texto reformulado:\n\n{texto}",
        "Registo académico":   "Adapta o seguinte texto para um registo académico em português, corrigindo erros e melhorando a precisão terminológica. Devolve apenas o texto:\n\n{texto}",
        "Resumo":              "Resume o seguinte texto em português mantendo as ideias principais. Devolve apenas o resumo:\n\n{texto}",
        "Correção gramatical": "Corrige os erros gramaticais e ortográficos do seguinte texto em português sem alterar o estilo. Devolve apenas o texto corrigido:\n\n{texto}",
    },
    "en": {
        "Normalização geral":  "Normalize the following English text by fixing spelling, punctuation and coherence. Return only the normalized text:\n\n{texto}",
        "Registo formal":      "Rewrite the following text in a formal professional English register. Return only the rewritten text:\n\n{texto}",
        "Registo académico":   "Adapt the following text to an academic English register. Return only the text:\n\n{texto}",
        "Resumo":              "Summarize the following English text keeping the main ideas. Return only the summary:\n\n{texto}",
        "Correção gramatical": "Fix all grammar and spelling errors in the following English text without changing its style. Return only the corrected text:\n\n{texto}",
    },
    "es": {
        "Normalização geral":  "Normaliza el siguiente texto en español corrigiendo errores. Devuelve solo el texto normalizado:\n\n{texto}",
        "Registo formal":      "Reescribe el siguiente texto en español en un registro formal. Devuelve solo el texto:\n\n{texto}",
        "Registo académico":   "Adapta el siguiente texto a un registro académico en español. Devuelve solo el texto:\n\n{texto}",
        "Resumo":              "Resume el siguiente texto en español. Devuelve solo el resumen:\n\n{texto}",
        "Correção gramatical": "Corrige los errores del siguiente texto en español. Devuelve solo el texto:\n\n{texto}",
    },
}


def gerar_prompt(texto: str, idioma: str, tipo: str) -> str:
    """
    Gera o prompt completo pronto a enviar à API.
    Seleciona o template certo com base no idioma detetado e no tipo de
    normalização escolhido pelo utilizador. Se o idioma não tiver templates
    definidos (ex: francês), usa os templates em inglês como fallback.
    Substitui {texto} no template pelo conteúdo real do chunk.
    """
    templates = TEMPLATES_PROMPT.get(idioma, TEMPLATES_PROMPT["en"])
    template = templates.get(tipo, list(templates.values())[0])
    return template.format(texto=texto)


# =============================================================================
# TAREFA 4 – API SLM
# =============================================================================

API_URL   = "https://reality.utad.net/slm"   # endpoint da API da UTAD
API_MODEL = "llama-3.2-1b-instruct"           # modelo a usar (definido no enunciado)


def chamar_api(prompt: str, timeout: int = 60, max_tentativas: int = 3) -> dict:
    """
    Envia um prompt para a API do SLM (Small Language Model) da UTAD.
    Usa apenas urllib da biblioteca padrão — sem bibliotecas externas.

    O corpo do pedido segue o formato OpenAI-compatible (igual ao ChatGPT):
    {
        "model": "llama-3.2-1b-instruct",
        "messages": [{"role": "user", "content": "..."}]
    }

    A resposta vem no mesmo formato:
    {"choices": [{"message": {"content": "resposta do modelo"}}]}

    Tem retry automático — se falhar, tenta de novo até max_tentativas vezes.
    Devolve um dicionário com:
    - sucesso: True/False
    - conteudo: texto gerado pelo modelo
    - tokens_prompt / tokens_resposta: uso de tokens
    - tempo: segundos totais
    - tentativas: quantas tentativas foram feitas (1, 2 ou 3)
    - erro: mensagem de erro se sucesso=False em todas as tentativas
    """
    # constrói o corpo do pedido HTTP em JSON — igual em todas as tentativas
    payload = {
        "model": API_MODEL,
        "messages": [{"role": "user", "content": prompt}]
    }
    corpo = json.dumps(payload).encode("utf-8")

    ultimo_erro = ""
    inicio_total = time.time()

    # tenta até max_tentativas vezes antes de desistir
    for tentativa in range(1, max_tentativas + 1):

        # cria o pedido HTTP de novo em cada tentativa
        pedido = urllib.request.Request(
            API_URL,
            data=corpo,
            method="POST",
            headers={"Content-Type": "application/json", "Accept": "application/json"},
        )

        inicio = time.time()
        try:
            with urllib.request.urlopen(pedido, timeout=timeout) as resp:
                tempo = round(time.time() - inicio, 2)
                dados = json.loads(resp.read().decode("utf-8"))

            # extrai o texto da resposta
            conteudo = ""
            if "choices" in dados and dados["choices"]:
                conteudo = dados["choices"][0].get("message", {}).get("content", "")
            elif "message" in dados:
                conteudo = dados["message"].get("content", "")

            uso = dados.get("usage", {})

            # sucesso — devolve com o número de tentativas que foram precisas
            return {
                "sucesso":         True,
                "conteudo":        conteudo,
                "tokens_prompt":   uso.get("prompt_tokens", 0),
                "tokens_resposta": uso.get("completion_tokens", 0),
                "tempo":           round(time.time() - inicio_total, 2),
                "tentativas":      tentativa,
                "erro":            None,
            }

        except urllib.error.HTTPError as e:
            # erro HTTP (ex: 500) — anota o erro e tenta de novo
            ultimo_erro = f"HTTP {e.code}: {e.reason}"

        except Exception as e:
            # timeout, sem ligação, etc. — anota e tenta de novo
            ultimo_erro = str(e)

        # se não foi a última tentativa, espera 2 segundos antes de tentar de novo
        # (evita sobrecarregar a API com pedidos imediatos)
        if tentativa < max_tentativas:
            time.sleep(2)

    # esgotou todas as tentativas sem sucesso
    return {
        "sucesso":         False,
        "conteudo":        "",
        "tokens_prompt":   0,
        "tokens_resposta": 0,
        "tempo":           round(time.time() - inicio_total, 2),
        "tentativas":      max_tentativas,
        "erro":            f"Falhou após {max_tentativas} tentativas. Último erro: {ultimo_erro}",
    }


# =============================================================================
# TAREFA 5 – GERAÇÃO DE RELATÓRIO HTML
# =============================================================================

def gerar_relatorio_html(nome_ficheiro, texto_bruto, texto_limpo,
                          estatisticas, idioma, confianca, estrategia,
                          tipo_norm, resultados_api=None, texto_final="") -> str:
    """
    Gera um relatório HTML completo com toda a informação do processo.
    O relatório inclui:
    - Visão geral com cards de estatísticas
    - Parâmetros usados (estratégia, tipo de normalização, idioma)
    - Tabela detalhada por etapa da pipeline
    - Texto original vs texto limpo (para comparação visual)
    - Resultados da API (se foi executada)
    Devolve uma string com o HTML completo pronto a guardar em ficheiro.
    """
    # função auxiliar para escapar caracteres HTML especiais
    # (< > & " ' podem partir o HTML se não forem escapados)
    def _e(t): return html_module.escape(str(t))

    agora = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # calcula a redução total de caracteres em toda a pipeline
    total_rem = sum(e["chars_removidos"] for e in estatisticas)
    pct = round(100 * total_rem / max(len(texto_bruto), 1), 1)

    # constrói as linhas da tabela de etapas
    linhas_etapas = ""
    for e in estatisticas:
        # verde se removeu chars (limpeza), vermelho se adicionou (raro)
        cor = "#4ade80" if e["chars_removidos"] >= 0 else "#f87171"
        linhas_etapas += (
            f"<tr><td>{_e(e['etapa'])}</td>"
            f"<td>{e['linhas_antes']:,}</td>"
            f"<td>{e['linhas_depois']:,}</td>"
            f"<td style='color:{cor};font-weight:600'>{e['chars_removidos']:+,}</td></tr>"
        )

    # secção dos resultados da API (só aparece se a API foi chamada)
    secao_api = ""
    if resultados_api:
        ok  = [r for r in resultados_api if r.get("sucesso")]
        avg = round(sum(r["tempo"] for r in resultados_api) / len(resultados_api), 2)
        toks = sum(r.get("tokens_resposta", 0) for r in resultados_api)
        amostra = _e(texto_final[:2000]) + ("..." if len(texto_final) > 2000 else "")
        secao_api = f"""<section class="s"><h2>4. Resultados da API SLM</h2>
        <div class="grid">
          <div class="card"><span class="val">{len(resultados_api)}</span><span class="lbl">Chunks enviados</span></div>
          <div class="card"><span class="val">{len(ok)}</span><span class="lbl">Com sucesso</span></div>
          <div class="card"><span class="val">{avg}s</span><span class="lbl">Tempo médio</span></div>
          <div class="card"><span class="val">{toks:,}</span><span class="lbl">Tokens gerados</span></div>
        </div>
        <h3>Texto normalizado (amostra)</h3><pre class="box">{amostra}</pre></section>"""

    # retorna o HTML completo com CSS inline para funcionar em qualquer browser
    return f"""<!DOCTYPE html><html lang="pt"><head><meta charset="UTF-8">
<title>TextNorm – Relatório</title><style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:Georgia,serif;background:#f8f7f4;color:#1a1a1a;line-height:1.7}}
.header{{background:#0d0d0d;color:#f0ede8;padding:2rem 3rem;border-bottom:3px solid #e85d26}}
.header h1{{font-family:'Courier New',monospace;font-size:1.6rem}}
.sub{{color:#888;font-style:italic;margin-top:.3rem}}
.meta{{margin-top:.8rem;font-size:.8rem;color:#666;font-family:monospace}}
.wrap{{max-width:900px;margin:0 auto;padding:2rem 1.5rem}}
.s{{background:#fff;border-radius:8px;padding:1.8rem;margin-bottom:1.2rem;border:1px solid #e8e4de}}
h2{{font-family:'Courier New',monospace;font-size:.95rem;color:#e85d26;text-transform:uppercase;
    letter-spacing:1px;border-bottom:1px solid #e8e4de;padding-bottom:.5rem;margin-bottom:1rem}}
h3{{font-size:.88rem;color:#555;margin:1rem 0 .3rem;font-family:'Courier New',monospace}}
.grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(120px,1fr));gap:.7rem;margin-bottom:1rem}}
.card{{background:#f8f7f4;border:1px solid #e8e4de;border-radius:6px;padding:.7rem;text-align:center}}
.val{{display:block;font-size:1.3rem;font-weight:700;font-family:'Courier New',monospace;color:#e85d26}}
.lbl{{display:block;font-size:.68rem;color:#888;text-transform:uppercase;letter-spacing:.5px;margin-top:.2rem}}
table{{width:100%;border-collapse:collapse;font-size:.83rem}}
th{{background:#f0ede8;padding:.5rem .8rem;text-align:left;font-family:'Courier New',monospace;
    font-size:.7rem;text-transform:uppercase;color:#666;border-bottom:2px solid #e8e4de}}
td{{padding:.5rem .8rem;border-bottom:1px solid #f0ede8}}
.box{{background:#f8f7f4;border:1px solid #e8e4de;border-radius:6px;padding:.9rem;
      font-family:'Courier New',monospace;font-size:.77rem;white-space:pre-wrap;
      word-break:break-word;max-height:250px;overflow-y:auto}}
.footer{{text-align:center;color:#aaa;font-size:.72rem;margin-top:2rem;font-family:monospace}}
</style></head><body>
<div class="header">
  <h1>⚙ TextNorm Pipeline — Relatório</h1>
  <div class="sub">Normalização de Texto com SLMs · UTAD 2025/26</div>
  <div class="meta">Ficheiro: {_e(nome_ficheiro)} &nbsp;|&nbsp; Gerado: {agora} &nbsp;|&nbsp; Modelo: {API_MODEL}</div>
</div>
<div class="wrap">
<section class="s"><h2>1. Visão Geral</h2>
  <div class="grid">
    <div class="card"><span class="val">{len(texto_bruto):,}</span><span class="lbl">Chars originais</span></div>
    <div class="card"><span class="val">{len(texto_limpo):,}</span><span class="lbl">Chars limpos</span></div>
    <div class="card"><span class="val">-{pct}%</span><span class="lbl">Redução</span></div>
    <div class="card"><span class="val">{_e(idioma).upper()}</span><span class="lbl">Idioma ({int(confianca*100)}%)</span></div>
    <div class="card"><span class="val">{len(estatisticas)}</span><span class="lbl">Etapas</span></div>
  </div>
  <table>
    <tr><th>Parâmetro</th><th>Valor</th></tr>
    <tr><td>Estratégia de chunking</td><td>{_e(estrategia)}</td></tr>
    <tr><td>Tipo de normalização</td><td>{_e(tipo_norm)}</td></tr>
    <tr><td>Idioma detetado</td><td>{_e(NOMES_IDIOMA.get(idioma, idioma))} ({int(confianca*100)}% confiança)</td></tr>
    <tr><td>Total caracteres removidos</td><td>{total_rem:,}</td></tr>
  </table>
</section>
<section class="s"><h2>2. Etapas da Pipeline</h2>
  <table><thead><tr><th>Etapa</th><th>Linhas antes</th><th>Linhas depois</th><th>Δ Chars</th></tr></thead>
  <tbody>{linhas_etapas}</tbody></table>
</section>
<section class="s"><h2>3. Texto Antes / Depois</h2>
  <h3>Texto original (primeiros 1500 chars)</h3>
  <pre class="box">{_e(texto_bruto[:1500])}{'...' if len(texto_bruto)>1500 else ''}</pre>
  <h3>Texto após limpeza (primeiros 1500 chars)</h3>
  <pre class="box">{_e(texto_limpo[:1500])}{'...' if len(texto_limpo)>1500 else ''}</pre>
</section>
{secao_api}
<div class="footer">TextNorm Pipeline · UTAD – Laboratório de Programação · {datetime.datetime.now().year}</div>
</div></body></html>"""


# =============================================================================
# INTERFACE GRÁFICA (tkinter)
# =============================================================================

# ── Paleta de cores e fontes ──────────────────────────────────────────────────
# Definidas como constantes para usar em toda a interface de forma consistente
COR_BG       = "#1a1a1a"   # fundo principal (cinzento muito escuro)
COR_PANEL    = "#242424"   # fundo dos painéis/cards (ligeiramente mais claro)
COR_BORDA    = "#333333"   # bordas dos elementos
COR_LARANJA  = "#E85D26"   # cor de destaque (botões, títulos, valores)
COR_TEXTO    = "#F0EDE8"   # texto principal (branco-creme)
COR_TEXTO2   = "#999999"   # texto secundário/subtítulos (cinzento)
COR_VERDE    = "#4ade80"   # indicador de sucesso/feito
COR_VERMELHO = "#f87171"   # indicador de erro/pendente
FONTE_MONO   = ("Courier New", 10)       # fonte monospace para código/texto
FONTE_TITULO = ("Courier New", 11, "bold")
FONTE_LABEL  = ("Segoe UI", 10)
FONTE_BTN    = ("Courier New", 10, "bold")


class TextNormApp(tk.Tk):
    """
    Classe principal da aplicação — herda de tk.Tk (a janela raiz do tkinter).
    Ao herdar de tk.Tk, esta classe É a janela principal da aplicação.
    Toda a interface, estado e lógica de interação estão aqui.
    """

    def __init__(self):
        super().__init__()   # inicializa a janela tkinter

        # configuração básica da janela
        self.title("⚙ TextNorm Pipeline — UTAD TP2")
        self.geometry("1100x750")    # tamanho inicial em píxeis
        self.minsize(900, 600)       # tamanho mínimo (não deixa encolher mais)
        self.configure(bg=COR_BG)

        # ── Estado da aplicação ───────────────────────────────────────────────
        # Estas variáveis guardam os dados entre separadores.
        # Quando o utilizador avança de separador, os dados ficam aqui.
        self.texto_bruto    = ""     # texto extraído do ficheiro (sem alterações)
        self.texto_limpo    = ""     # texto após pipeline de limpeza
        self.nome_ficheiro  = ""     # nome do ficheiro carregado
        self.idioma         = "desconhecido"  # idioma detetado
        self.confianca_id   = 0.0    # confiança da deteção de idioma (0 a 1)
        self.estatisticas   = []     # estatísticas de cada etapa da pipeline
        self.chunks         = []     # lista de chunks gerados
        self.prompts        = []     # lista de prompts (um por chunk)
        self.resultados_api = []     # respostas da API (uma por chunk enviado)
        self.texto_final    = ""     # texto normalizado final (juntando todas as respostas)

        # variáveis tkinter para os controlos da interface (ligadas a widgets)
        self.etapas_vars    = []     # lista de BooleanVar para os checkboxes da pipeline
        self.estrategia_var = tk.StringVar(value="Por parágrafo")   # dropdown do chunking
        self.tipo_norm_var  = tk.StringVar(value="Normalização geral")  # dropdown do tipo
        self.max_chunks_var = tk.IntVar(value=3)   # spinbox do máx de chunks para API

        self._build_ui()   # constrói toda a interface gráfica

    # ── Métodos auxiliares para criar widgets ─────────────────────────────────
    # Estes métodos são atalhos para criar widgets sempre com as mesmas cores/fontes,
    # evitando repetir os mesmos argumentos em todos os sítios

    def _frame(self, parent, **kw):
        """Cria um Frame com a cor de fundo padrão."""
        return tk.Frame(parent, bg=COR_BG, **kw)

    def _panel(self, parent, titulo="", **kw):
        """Cria um painel com borda e título opcional."""
        outer = tk.Frame(parent, bg=COR_PANEL, bd=0, highlightthickness=1,
                         highlightbackground=COR_BORDA, **kw)
        if titulo:
            tk.Label(outer, text=titulo.upper(), font=("Courier New", 8, "bold"),
                     bg=COR_PANEL, fg=COR_LARANJA).pack(anchor="w", padx=10, pady=(8,2))
        return outer

    def _label(self, parent, texto, **kw):
        """Cria um Label com o estilo padrão da aplicação."""
        return tk.Label(parent, text=texto, bg=kw.pop("bg", COR_BG),
                        fg=kw.pop("fg", COR_TEXTO), font=FONTE_LABEL, **kw)

    def _btn(self, parent, texto, comando, cor=COR_LARANJA, **kw):
        """Cria um botão com o estilo padrão."""
        return tk.Button(parent, text=texto, command=comando,
                         bg=cor, fg="white", font=FONTE_BTN,
                         relief="flat", cursor="hand2",
                         activebackground="#c94d1e", activeforeground="white",
                         padx=14, pady=6, **kw)

    def _textarea(self, parent, height=10, **kw):
        """Cria uma área de texto com scroll integrado e estilo escuro."""
        return scrolledtext.ScrolledText(
            parent, height=height, font=FONTE_MONO,
            bg="#0f0f0f", fg=COR_TEXTO, insertbackground=COR_TEXTO,
            relief="flat", wrap="word", bd=0,
            selectbackground=COR_LARANJA, **kw)

    def _status(self, msg):
        """Atualiza a mensagem na barra de estado no fundo da janela."""
        self.status_var.set(msg)
        self.update_idletasks()   # força o tkinter a redesenhar imediatamente

    def _stat_card(self, parent, titulo, valor):
        """Cria um card de estatística (valor grande + label pequeno por baixo)."""
        f = tk.Frame(parent, bg=COR_PANEL, bd=0, highlightthickness=1,
                     highlightbackground=COR_BORDA)
        f.pack(side="left", padx=4, pady=4, fill="x", expand=True)
        tk.Label(f, text=str(valor), font=("Courier New", 14, "bold"),
                 bg=COR_PANEL, fg=COR_LARANJA).pack(pady=(8,0))
        tk.Label(f, text=titulo.upper(), font=("Courier New", 7),
                 bg=COR_PANEL, fg=COR_TEXTO2).pack(pady=(0,8))
        return f

    # ── Construção da UI principal ────────────────────────────────────────────

    def _build_ui(self):
        """
        Constrói toda a estrutura visual da aplicação:
        1. Barra de cabeçalho no topo
        2. Notebook (sistema de separadores) com 5 tabs
        3. Barra de estado no fundo
        """
        # Cabeçalho superior
        header = tk.Frame(self, bg="#0d0d0d", height=55)
        header.pack(fill="x")
        header.pack_propagate(False)   # impede que o frame encolha para caber o conteúdo
        tk.Label(header, text="⚙  TextNorm Pipeline",
                 font=("Courier New", 14, "bold"), bg="#0d0d0d", fg=COR_TEXTO).pack(side="left", padx=20, pady=10)
        tk.Label(header, text="UTAD · Laboratório de Programação · TP2 · 2025/26",
                 font=("Courier New", 9), bg="#0d0d0d", fg=COR_TEXTO2).pack(side="left", padx=5, pady=10)
        # linha laranja decorativa abaixo do cabeçalho
        tk.Frame(self, bg=COR_LARANJA, height=3).pack(fill="x")

        # Configura o estilo visual dos separadores (tabs)
        style = ttk.Style(self)
        style.theme_use("default")
        style.configure("TNotebook", background=COR_BG, borderwidth=0)
        style.configure("TNotebook.Tab", background=COR_PANEL, foreground=COR_TEXTO2,
                        font=FONTE_TITULO, padding=[16, 8], borderwidth=0)
        style.map("TNotebook.Tab",
                  background=[("selected", COR_LARANJA)],    # tab ativa fica laranja
                  foreground=[("selected", "white")])
        style.configure("TFrame", background=COR_BG)

        # Cria o Notebook (contentor de separadores)
        self.nb = ttk.Notebook(self)
        self.nb.pack(fill="both", expand=True)

        # Cria os 5 frames (um por separador)
        self.tab1 = ttk.Frame(self.nb)
        self.tab2 = ttk.Frame(self.nb)
        self.tab3 = ttk.Frame(self.nb)
        self.tab4 = ttk.Frame(self.nb)
        self.tab5 = ttk.Frame(self.nb)

        # Adiciona cada frame ao notebook com o seu nome
        self.nb.add(self.tab1, text="  1 · Extração  ")
        self.nb.add(self.tab2, text="  2 · Limpeza  ")
        self.nb.add(self.tab3, text="  3 · Chunking & Prompt  ")
        self.nb.add(self.tab4, text="  4 · API SLM  ")
        self.nb.add(self.tab5, text="  5 · Relatório  ")

        # Constrói o conteúdo de cada separador
        self._build_tab1()
        self._build_tab2()
        self._build_tab3()
        self._build_tab4()
        self._build_tab5()

        # Barra de estado no fundo — mostra mensagens de progresso
        self.status_var = tk.StringVar(value="Pronto. Carrega um ficheiro para começar.")
        barra = tk.Frame(self, bg="#0d0d0d", height=28)
        barra.pack(fill="x", side="bottom")
        tk.Label(barra, textvariable=self.status_var,
                 font=("Courier New", 9), bg="#0d0d0d", fg=COR_TEXTO2,
                 anchor="w").pack(fill="x", padx=12, pady=4)

    # ── Separador 1 – Extração ────────────────────────────────────────────────

    def _build_tab1(self):
        """Constrói o separador de extração de texto."""
        p = self._frame(self.tab1)
        p.pack(fill="both", expand=True, padx=16, pady=14)

        self._label(p, "Carrega um ficheiro PDF, DOCX ou TXT. O texto é apresentado em bruto, sem qualquer transformação.",
                    fg=COR_TEXTO2).pack(anchor="w", pady=(0,10))

        # Linha do botão de upload + nome do ficheiro
        top = self._frame(p)
        top.pack(fill="x")
        self._btn(top, " Abrir Ficheiro", self._abrir_ficheiro).pack(side="left")
        self.label_ficheiro = tk.Label(top, text="Nenhum ficheiro selecionado",
                                       font=("Courier New", 9), bg=COR_BG, fg=COR_TEXTO2)
        self.label_ficheiro.pack(side="left", padx=14)

        # Frame para os cards de estatísticas (preenchido depois da extração)
        self.stats_frame1 = self._frame(p)
        self.stats_frame1.pack(fill="x", pady=(10,6))

        # Label para mostrar o idioma detetado
        self.label_idioma = tk.Label(p, text="", font=("Courier New", 9),
                                     bg=COR_BG, fg=COR_VERDE)
        self.label_idioma.pack(anchor="w", pady=(0,6))

        self._label(p, "Texto extraído (bruto):").pack(anchor="w")
        # área de texto grande para mostrar o conteúdo extraído
        self.area_bruto = self._textarea(p, height=22)
        self.area_bruto.pack(fill="both", expand=True, pady=(4,0))

    def _abrir_ficheiro(self):
        """
        Chamada quando o utilizador clica em 'Abrir Ficheiro'.
        Abre uma janela de diálogo do sistema operativo para selecionar o ficheiro.
        Depois chama extrair_texto() e atualiza toda a interface com os resultados.
        Em caso de erro mostra uma janela de alerta.
        """
        # filedialog.askopenfilename abre a janela nativa de 'Abrir ficheiro'
        caminho = filedialog.askopenfilename(
            title="Seleciona um ficheiro",
            filetypes=[("Documentos", "*.pdf *.docx *.txt"), ("Todos", "*.*")]
        )
        if not caminho:
            return   # utilizador cancelou — não faz nada

        self._status(f"A extrair {os.path.basename(caminho)}...")
        try:
            texto = extrair_texto(caminho)
            self.texto_bruto   = texto
            self.nome_ficheiro = os.path.basename(caminho)

            # reset de todo o estado downstream (limpeza em diante)
            # para evitar que dados de um ficheiro anterior fiquem misturados
            self.texto_limpo = ""
            self.estatisticas = []
            self.chunks = []
            self.prompts = []
            self.resultados_api = []
            self.texto_final = ""

            # deteta idioma e guarda no estado
            idioma, conf = detetar_idioma(texto)
            self.idioma = idioma
            self.confianca_id = conf

            # atualiza a interface com os resultados
            self.label_ficheiro.config(text=self.nome_ficheiro)
            self._atualizar_stats1(texto)
            self.label_idioma.config(
                text=f" Idioma detetado: {NOMES_IDIOMA.get(idioma, idioma)}  (confiança {int(conf*100)}%)")

            # coloca o texto na área — config(state="normal") permite edição,
            # depois de inserir voltamos a "disabled" para o utilizador não editar
            self.area_bruto.config(state="normal")
            self.area_bruto.delete("1.0", "end")   # limpa o conteúdo anterior
            self.area_bruto.insert("1.0", texto)   # insere o novo texto
            self.area_bruto.config(state="disabled")

            self._status(f"✓ Ficheiro carregado — {len(texto):,} caracteres")
        except Exception as e:
            messagebox.showerror("Erro na extração", str(e))
            self._status("Erro na extração.")

    def _atualizar_stats1(self, texto):
        """Atualiza os cards de estatísticas do separador 1."""
        # limpa os cards anteriores (se existirem)
        for w in self.stats_frame1.winfo_children():
            w.destroy()
        self._stat_card(self.stats_frame1, "Caracteres", f"{len(texto):,}")
        self._stat_card(self.stats_frame1, "Palavras",   f"{len(texto.split()):,}")
        self._stat_card(self.stats_frame1, "Linhas",     f"{texto.count(chr(10))+1:,}")
        self._stat_card(self.stats_frame1, "Tamanho",    f"{len(texto.encode())/1024:.1f} KB")

    # ── Separador 2 – Limpeza ─────────────────────────────────────────────────

    def _build_tab2(self):
        """Constrói o separador de limpeza com checkboxes para cada etapa."""
        p = self._frame(self.tab2)
        p.pack(fill="both", expand=True, padx=16, pady=14)

        # Painel com os checkboxes das etapas
        painel = self._panel(p, "Etapas da pipeline")
        painel.pack(fill="x", pady=(0,10))
        grid = tk.Frame(painel, bg=COR_PANEL)
        grid.pack(fill="x", padx=10, pady=(0,10))

        # cria um checkbox para cada etapa da pipeline
        # cada checkbox tem uma BooleanVar associada — True = ativa, False = desativada
        self.etapas_vars = []
        for i, (nome, _) in enumerate(ETAPAS_PIPELINE):
            var = tk.BooleanVar(value=True)   # todas ativas por defeito
            cb = tk.Checkbutton(grid, text=nome, variable=var,
                                bg=COR_PANEL, fg=COR_TEXTO, selectcolor=COR_BG,
                                activebackground=COR_PANEL, activeforeground=COR_TEXTO,
                                font=FONTE_LABEL, cursor="hand2")
            # organiza em 2 colunas (i//2 = linha, i%2 = coluna)
            cb.grid(row=i//2, column=i%2, sticky="w", padx=10, pady=2)
            self.etapas_vars.append(var)

        self._btn(p, " Executar Pipeline", self._correr_pipeline).pack(anchor="w", pady=(0,10))

        # Frame para cards de estatísticas (preenchido após execução)
        self.stats_frame2 = self._frame(p)
        self.stats_frame2.pack(fill="x", pady=(0,6))

        # Visualização lado a lado: antes vs depois
        colunas = self._frame(p)
        colunas.pack(fill="both", expand=True)

        col_e = self._frame(colunas)
        col_e.pack(side="left", fill="both", expand=True, padx=(0,6))
        self._label(col_e, "Antes da limpeza:").pack(anchor="w")
        self.area_antes = self._textarea(col_e, height=16)
        self.area_antes.pack(fill="both", expand=True, pady=(4,0))

        col_d = self._frame(colunas)
        col_d.pack(side="left", fill="both", expand=True)
        self._label(col_d, "Depois da limpeza:").pack(anchor="w")
        self.area_depois = self._textarea(col_d, height=16)
        self.area_depois.pack(fill="both", expand=True, pady=(4,0))

    def _correr_pipeline(self):
        """
        Chamada quando o utilizador clica em 'Executar Pipeline'.
        Lê quais os checkboxes que estão ativos, chama correr_pipeline()
        e atualiza a interface com os resultados e estatísticas.
        """
        if not self.texto_bruto:
            messagebox.showwarning("Aviso", "Extrai primeiro um ficheiro no separador 1.")
            return

        # lê o estado de cada checkbox (True/False)
        ativas = [v.get() for v in self.etapas_vars]

        self._status("A processar pipeline...")
        try:
            limpo, stats = correr_pipeline(self.texto_bruto, ativas)
            self.texto_limpo  = limpo
            self.estatisticas = stats

            # reset downstream
            self.chunks = []
            self.prompts = []
            self.resultados_api = []
            self.texto_final = ""

            # atualiza os cards de estatísticas
            for w in self.stats_frame2.winfo_children():
                w.destroy()
            total_rem = sum(e["chars_removidos"] for e in stats)
            pct = round(100 * total_rem / max(len(self.texto_bruto), 1), 1)
            self._stat_card(self.stats_frame2, "Etapas",           len(stats))
            self._stat_card(self.stats_frame2, "Chars removidos",  f"{total_rem:,}")
            self._stat_card(self.stats_frame2, "Redução",          f"{pct}%")
            self._stat_card(self.stats_frame2, "Chars finais",     f"{len(limpo):,}")

            # mostra o texto antes e depois lado a lado
            for area, conteudo in [(self.area_antes, self.texto_bruto),
                                   (self.area_depois, limpo)]:
                area.config(state="normal")
                area.delete("1.0", "end")
                area.insert("1.0", conteudo)
                area.config(state="disabled")

            self._status(f"✓ Pipeline concluída — {len(stats)} etapas | -{pct}% de redução")
        except Exception as e:
            messagebox.showerror("Erro", str(e))
            self._status("Erro na pipeline.")

    # ── Separador 3 – Chunking & Prompt ──────────────────────────────────────

    def _build_tab3(self):
        """Constrói o separador de chunking e geração de prompts."""
        p = self._frame(self.tab3)
        p.pack(fill="both", expand=True, padx=16, pady=14)

        # Painel de opções com dois dropdowns
        opcoes = self._panel(p, "Configuração")
        opcoes.pack(fill="x", pady=(0,10))
        grid = tk.Frame(opcoes, bg=COR_PANEL)
        grid.pack(fill="x", padx=10, pady=(0,10))

        self._label(grid, "Estratégia de chunking:", bg=COR_PANEL).grid(row=0, column=0, sticky="w", pady=4, padx=4)
        # Combobox = dropdown de seleção
        ttk.Combobox(grid, textvariable=self.estrategia_var,
                     values=list(ESTRATEGIAS_CHUNKING.keys()),
                     state="readonly", width=28).grid(row=0, column=1, sticky="w", pady=4, padx=4)

        self._label(grid, "Tipo de normalização:", bg=COR_PANEL).grid(row=1, column=0, sticky="w", pady=4, padx=4)
        tipos = list(TEMPLATES_PROMPT["pt"].keys())
        ttk.Combobox(grid, textvariable=self.tipo_norm_var,
                     values=tipos, state="readonly", width=28).grid(row=1, column=1, sticky="w", pady=4, padx=4)

        self._btn(p, " Segmentar & Gerar Prompts", self._correr_chunking).pack(anchor="w", pady=(0,10))

        # Cards de estatísticas dos chunks
        self.stats_frame3 = self._frame(p)
        self.stats_frame3.pack(fill="x", pady=(0,6))

        # Dropdown para selecionar qual chunk visualizar
        sel_frame = self._frame(p)
        sel_frame.pack(fill="x", pady=(0,6))
        self._label(sel_frame, "Ver chunk:").pack(side="left")
        self.chunk_var = tk.StringVar()
        self.combo_chunks = ttk.Combobox(sel_frame, textvariable=self.chunk_var,
                                          state="readonly", width=30)
        self.combo_chunks.pack(side="left", padx=8)
        # quando o utilizador escolhe um chunk no dropdown, chama _mostrar_chunk
        self.combo_chunks.bind("<<ComboboxSelected>>", self._mostrar_chunk)

        # Visualização lado a lado: chunk vs prompt
        colunas = self._frame(p)
        colunas.pack(fill="both", expand=True)

        col_c = self._frame(colunas)
        col_c.pack(side="left", fill="both", expand=True, padx=(0,6))
        self._label(col_c, "Texto do chunk:").pack(anchor="w")
        self.area_chunk = self._textarea(col_c, height=16)
        self.area_chunk.pack(fill="both", expand=True, pady=(4,0))

        col_pr = self._frame(colunas)
        col_pr.pack(side="left", fill="both", expand=True)
        self._label(col_pr, "Prompt gerado:").pack(anchor="w")
        self.area_prompt = self._textarea(col_pr, height=16)
        self.area_prompt.pack(fill="both", expand=True, pady=(4,0))

    def _correr_chunking(self):
        """
        Chamada quando o utilizador clica em 'Segmentar & Gerar Prompts'.
        Usa o texto limpo (se disponível) ou o bruto, divide em chunks
        com a estratégia selecionada e gera um prompt para cada chunk.
        """
        # usa o texto limpo se a pipeline foi executada, senão usa o bruto
        fonte = self.texto_limpo or self.texto_bruto
        if not fonte:
            messagebox.showwarning("Aviso", "Extrai primeiro um ficheiro no separador 1.")
            return

        estrategia = self.estrategia_var.get()
        tipo = self.tipo_norm_var.get()

        # obtém a função de chunking correspondente à estratégia selecionada
        fn = ESTRATEGIAS_CHUNKING.get(estrategia, list(ESTRATEGIAS_CHUNKING.values())[0])
        self.chunks  = fn(fonte)   # divide o texto em chunks
        # para cada chunk, gera o prompt no idioma detetado e tipo escolhido
        self.prompts = [gerar_prompt(c, self.idioma, tipo) for c in self.chunks]

        # reset downstream
        self.resultados_api = []
        self.texto_final = ""

        # atualiza os cards de estatísticas
        for w in self.stats_frame3.winfo_children():
            w.destroy()
        avg = sum(len(c.split()) for c in self.chunks) / max(len(self.chunks), 1)
        self._stat_card(self.stats_frame3, "Chunks",        len(self.chunks))
        self._stat_card(self.stats_frame3, "Palavras/chunk", f"{avg:.0f}")
        self._stat_card(self.stats_frame3, "Maior",         f"{max(len(c) for c in self.chunks):,} chars")
        self._stat_card(self.stats_frame3, "Menor",         f"{min(len(c) for c in self.chunks):,} chars")

        # preenche o dropdown de seleção de chunks
        self.combo_chunks["values"] = [
            f"Chunk {i+1}  ({len(c.split())} palavras)"
            for i, c in enumerate(self.chunks)
        ]
        if self.chunks:
            self.combo_chunks.current(0)   # seleciona o primeiro por defeito
            self._mostrar_chunk()           # mostra o conteúdo do primeiro chunk

        self._status(f"✓ {len(self.chunks)} chunks gerados com estratégia '{estrategia}'")

    def _mostrar_chunk(self, event=None):
        """
        Atualiza as áreas de texto do separador 3 com o chunk e prompt selecionados.
        Chamada automaticamente quando o utilizador muda a seleção no dropdown.
        """
        idx = self.combo_chunks.current()   # índice do item selecionado
        if idx < 0 or idx >= len(self.chunks):
            return
        # atualiza as duas áreas de texto com o chunk e prompt correspondentes
        for area, conteudo in [(self.area_chunk, self.chunks[idx]),
                               (self.area_prompt, self.prompts[idx])]:
            area.config(state="normal")
            area.delete("1.0", "end")
            area.insert("1.0", conteudo)
            area.config(state="disabled")

    # ── Separador 4 – API SLM ─────────────────────────────────────────────────

    def _build_tab4(self):
        """Constrói o separador de comunicação com a API do SLM."""
        p = self._frame(self.tab4)
        p.pack(fill="both", expand=True, padx=16, pady=14)

        # painel informativo com o URL e modelo
        info = self._panel(p, "Endpoint")
        info.pack(fill="x", pady=(0,10))
        tk.Label(info, text=f"URL: {API_URL}    Modelo: {API_MODEL}",
                 font=FONTE_MONO, bg=COR_PANEL, fg=COR_TEXTO2).pack(anchor="w", padx=10, pady=6)

        # controlo do número máximo de chunks a enviar
        cfg = self._frame(p)
        cfg.pack(fill="x", pady=(0,10))
        self._label(cfg, "Máx. chunks a enviar:").pack(side="left")
        # Spinbox = campo numérico com botões +/-
        tk.Spinbox(cfg, from_=1, to=20, textvariable=self.max_chunks_var,
                   width=4, bg=COR_PANEL, fg=COR_TEXTO, font=FONTE_MONO,
                   buttonbackground=COR_PANEL).pack(side="left", padx=8)

        btns = self._frame(p)
        btns.pack(fill="x", pady=(0,10))
        self._btn(btns, " Enviar para API", self._enviar_api).pack(side="left", padx=(0,8))

        # barra de progresso — atualizada durante o envio dos chunks
        self.progresso_var = tk.DoubleVar(value=0)
        self.barra_prog = ttk.Progressbar(p, variable=self.progresso_var, maximum=100)
        self.barra_prog.pack(fill="x", pady=(0,8))

        # cards de estatísticas (preenchidos após receber respostas)
        self.stats_frame4 = self._frame(p)
        self.stats_frame4.pack(fill="x", pady=(0,6))

        self._label(p, "Texto normalizado pela API:").pack(anchor="w")
        self.area_normalizado = self._textarea(p, height=18)
        self.area_normalizado.pack(fill="both", expand=True, pady=(4,0))

    def _enviar_api(self):
        """
        Chamada quando o utilizador clica em 'Enviar para API'.
        Usa threading para executar os pedidos HTTP em segundo plano —
        sem isso, a interface ficaria completamente bloqueada/congelada
        durante o tempo que a API demora a responder (pode ser vários segundos).
        O threading.Thread cria uma thread separada que corre a função 'tarefa'.
        No final, usa self.after(0, ...) para atualizar a interface na thread principal
        (tkinter não é thread-safe — só a thread principal pode modificar widgets).
        """
        if not self.prompts:
            messagebox.showwarning("Aviso", "Gera os prompts no separador 3 primeiro.")
            return

        n = min(self.max_chunks_var.get(), len(self.prompts))
        self._status(f"A enviar {n} chunks para a API...")
        self.progresso_var.set(0)

        def tarefa():
            """Função que corre em segundo plano — faz os pedidos à API com retry automático."""
            resultados = []
            for i, prompt in enumerate(self.prompts[:n]):
                # informa que está a tentar pela primeira vez
                self._status(f"Chunk {i+1}/{n} — Tentativa 1...")
                # chamar_api tenta automaticamente até 3 vezes se falhar
                r = chamar_api(prompt)
                resultados.append(r)
                self.progresso_var.set((i + 1) / n * 100)
                # mostra o resultado com o número de tentativas que foram precisas
                tentativas = r.get("tentativas", 1)
                if r["sucesso"]:
                    if tentativas == 1:
                        self._status(f"Chunk {i+1}/{n} — OK à 1ª tentativa ({r['tempo']}s)")
                    else:
                        self._status(f"Chunk {i+1}/{n} — OK à {tentativas}ª tentativa ({r['tempo']}s)")
                else:
                    self._status(f"Chunk {i+1}/{n} — Erro após {tentativas} tentativas: {r['erro']}")

            self.resultados_api = resultados
            self.texto_final = "\n\n".join(
                r["conteudo"] for r in resultados if r.get("sucesso") and r.get("conteudo"))
            self.after(0, self._atualizar_tab4)

        # daemon=True → a thread termina automaticamente quando a aplicação fecha
        threading.Thread(target=tarefa, daemon=True).start()

    def _atualizar_tab4(self):
        """
        Atualiza a interface após receber todas as respostas da API.
        Chamada na thread principal via self.after(0, ...) para ser segura.
        """
        resultados = self.resultados_api
        ok   = sum(1 for r in resultados if r.get("sucesso"))
        fail = len(resultados) - ok
        toks = sum(r.get("tokens_resposta", 0) for r in resultados)
        avg  = round(sum(r["tempo"] for r in resultados) / max(len(resultados), 1), 2)

        # limpa e reconstrói os cards de estatísticas
        for w in self.stats_frame4.winfo_children():
            w.destroy()
        # tentativas totais feitas em todos os chunks
        total_tentativas = sum(r.get("tentativas", 1) for r in resultados)
        self._stat_card(self.stats_frame4, "Enviados",    len(resultados))
        self._stat_card(self.stats_frame4, "Sucesso",     ok)
        self._stat_card(self.stats_frame4, "Erro",        fail)
        self._stat_card(self.stats_frame4, "Tentativas",  total_tentativas)
        self._stat_card(self.stats_frame4, "Tokens",      f"{toks:,}")
        self._stat_card(self.stats_frame4, "Tempo médio", f"{avg}s")

        # mostra o texto normalizado final
        self.area_normalizado.config(state="normal")
        self.area_normalizado.delete("1.0", "end")
        self.area_normalizado.insert("1.0", self.texto_final or "(Sem resposta da API)")
        self.area_normalizado.config(state="disabled")

        self._status(f"✓ {ok}/{len(resultados)} chunks processados | {toks} tokens gerados")

    # ── Separador 5 – Relatório ───────────────────────────────────────────────

    def _build_tab5(self):
        """Constrói o separador de geração de relatórios."""
        p = self._frame(self.tab5)
        p.pack(fill="both", expand=True, padx=16, pady=14)

        self._label(p, "Gera um relatório completo com os parâmetros, o texto antes/depois e a avaliação da normalização.",
                    fg=COR_TEXTO2).pack(anchor="w", pady=(0,14))

        # Painel de estado — mostra quais as etapas já feitas
        estado = self._panel(p, "Estado do pipeline")
        estado.pack(fill="x", pady=(0,14))
        self.estado_frame = tk.Frame(estado, bg=COR_PANEL)
        self.estado_frame.pack(fill="x", padx=10, pady=(0,10))
        self._atualizar_estado()

        # Botões de exportação
        btns = self._frame(p)
        btns.pack(fill="x", pady=(0,14))
        self._btn(btns, " Exportar HTML", lambda: self._gerar_relatorio("html")).pack(side="left", padx=(0,10))
        self._btn(btns, " Exportar PDF",  lambda: self._gerar_relatorio("pdf"),
                  cor="#2563eb").pack(side="left")
        tk.Label(btns, text="(PDF requer weasyprint)", font=("Courier New", 8),
                 bg=COR_BG, fg=COR_TEXTO2).pack(side="left", padx=10)

        # Área de pré-visualização do HTML gerado
        self._label(p, "Pré-visualização do relatório:").pack(anchor="w")
        self.area_relatorio = self._textarea(p, height=22)
        self.area_relatorio.pack(fill="both", expand=True, pady=(4,0))

        # quando o utilizador muda para este separador, atualiza o estado
        self.nb.bind("<<NotebookTabChanged>>", self._on_tab_change)

    def _on_tab_change(self, event):
        """Chamada pelo tkinter sempre que o utilizador muda de separador."""
        if self.nb.index("current") == 4:   # índice 4 = separador 5 (começa em 0)
            self._atualizar_estado()

    def _atualizar_estado(self):
        """
        Atualiza os indicadores de estado no separador 5.
        Para cada etapa, verifica se os dados correspondentes existem no estado
        e mostra ✓ Feito (verde) ou Pendente (vermelho).
        """
        for w in self.estado_frame.winfo_children():
            w.destroy()   # limpa os indicadores anteriores

        # lista de (nome, condição) — a condição é True se a etapa foi feita
        itens = [
            (" Extração",  bool(self.texto_bruto)),
            (" Limpeza",   bool(self.texto_limpo)),
            (" Chunking",  bool(self.chunks)),
            (" API SLM",   bool(self.resultados_api)),
        ]
        for nome, feito in itens:
            cor  = COR_VERDE if feito else COR_VERMELHO
            tick = "✓ Feito" if feito else "⏳ Pendente"
            f = tk.Frame(self.estado_frame, bg=COR_PANEL)
            f.pack(side="left", expand=True, fill="x", padx=6, pady=6)
            tk.Label(f, text=nome, font=("Courier New", 9), bg=COR_PANEL, fg=COR_TEXTO).pack()
            tk.Label(f, text=tick, font=("Courier New", 9, "bold"), bg=COR_PANEL, fg=cor).pack()

    def _gerar_relatorio(self, fmt):
        """
        Gera o relatório e abre a janela 'Guardar como' para o utilizador
        escolher onde guardar.
        fmt = "html" → guarda como ficheiro .html
        fmt = "pdf"  → converte para PDF com weasyprint e guarda como .pdf
                       Se weasyprint não estiver instalado, cai back para HTML.
        """
        if not self.texto_bruto:
            messagebox.showwarning("Aviso", "Extrai primeiro um ficheiro no separador 1.")
            return
        try:
            # gera o HTML do relatório com todos os dados do estado atual
            html = gerar_relatorio_html(
                nome_ficheiro  = self.nome_ficheiro or "texto",
                texto_bruto    = self.texto_bruto,
                texto_limpo    = self.texto_limpo or self.texto_bruto,
                estatisticas   = self.estatisticas,
                idioma         = self.idioma,
                confianca      = self.confianca_id,
                estrategia     = self.estrategia_var.get(),
                tipo_norm      = self.tipo_norm_var.get(),
                resultados_api = self.resultados_api or None,
                texto_final    = self.texto_final,
            )

            if fmt == "html":
                # abre janela "Guardar como" do sistema operativo
                caminho = filedialog.asksaveasfilename(
                    defaultextension=".html",
                    filetypes=[("HTML", "*.html")],
                    initialfile="relatorio_textnorm.html",
                )
                if caminho:
                    with open(caminho, "w", encoding="utf-8") as f:
                        f.write(html)
                    self._status(f"✓ Relatório HTML guardado em {caminho}")
                    messagebox.showinfo("Sucesso", f"Relatório guardado em:\n{caminho}")

            elif fmt == "pdf":
                try:
                    from weasyprint import HTML as WP   # importação opcional
                    caminho = filedialog.asksaveasfilename(
                        defaultextension=".pdf",
                        filetypes=[("PDF", "*.pdf")],
                        initialfile="relatorio_textnorm.pdf",
                    )
                    if caminho:
                        WP(string=html).write_pdf(caminho)   # converte HTML → PDF
                        self._status(f"✓ Relatório PDF guardado em {caminho}")
                        messagebox.showinfo("Sucesso", f"Relatório PDF guardado em:\n{caminho}")
                except ImportError:
                    # weasyprint não instalado — avisa e guarda como HTML
                    messagebox.showerror("weasyprint não encontrado",
                                         "Instala com:\n\npip install weasyprint\n\nA guardar como HTML em alternativa.")
                    self._gerar_relatorio("html")
                    return

            # mostra o HTML gerado na área de pré-visualização
            self.area_relatorio.config(state="normal")
            self.area_relatorio.delete("1.0", "end")
            self.area_relatorio.insert("1.0", html)
            self.area_relatorio.config(state="disabled")

        except Exception as e:
            messagebox.showerror("Erro", str(e))
            self._status("Erro ao gerar relatório.")


# =============================================================================
# PONTO DE ENTRADA DO PROGRAMA
# =============================================================================

if __name__ == "__main__":
    # __name__ == "__main__" garante que este código só corre quando o ficheiro é executado diretamente (python textnorm_gui.py) e não quando é importado
    app = TextNormApp()   # cria a janela principal
    app.mainloop()        # inicia o loop de eventos do tkinter — mantém a janela aberta e processa cliques, teclado, etc. até a janela ser fechada